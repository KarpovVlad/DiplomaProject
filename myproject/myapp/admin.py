from datetime import datetime
from django.urls import reverse
from import_export.admin import ImportExportModelAdmin
from import_export.results import RowResult
from .models import University, Faculty, Department, Professor, Course, Student, CourseApplication, CourseEnrollment
from docx import Document
from django.utils.html import format_html
from django.core.management import call_command
from django.contrib import admin
from django.contrib.auth.admin import UserAdmin
from .forms import BulkUserCreationForm
from django.shortcuts import render, redirect
from django.urls import path
from django.http import HttpResponse
from django.utils.text import slugify
from import_export import resources, fields
from import_export.widgets import ForeignKeyWidget
from django.contrib.auth.models import User
from django.db import models
import io
import csv
import random
import string

temporary_passwords = {}


def generate_idp_for_student(student_id):
    student = Student.objects.get(id=student_id)
    enrollments = CourseEnrollment.objects.filter(student=student)

    document = Document()
    document.add_heading('Індивідуальний навчальний план', level=1)

    document.add_paragraph(f"Студент: {student.user.first_name} {student.user.last_name}")
    document.add_paragraph(f"Кафедра: {student.department.name}")
    document.add_paragraph(f"Середній бал: {student.average_grade}")

    document.add_heading('Обрані дисципліни:', level=2)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Назва дисципліни'
    hdr_cells[1].text = 'Викладач'

    for enrollment in enrollments:
        row_cells = table.add_row().cells
        row_cells[0].text = enrollment.course.name
        row_cells[1].text = str(enrollment.course.professor)

    document.add_page_break()

    file_path = f'documents/{student.user.username}_IDP.docx'
    document.save(file_path)

    student.document = file_path
    student.save()


def generate_idp_for_group(department_id):
    students = Student.objects.filter(department_id=department_id)

    for student in students:
        generate_idp_for_student(student.id)


class UniversityAdmin(ImportExportModelAdmin):
    list_display = ('name',)
    search_fields = ('name',)


class FacultyAdmin(ImportExportModelAdmin):
    list_display = ('name', 'university')
    search_fields = ('name', 'university__name')
    list_filter = ('university__name',)


class DepartmentAdmin(ImportExportModelAdmin):
    list_display = ('name', 'faculty')
    search_fields = ('name', 'faculty__name')
    list_filter = ('faculty__name', 'faculty__university__name')


class ProfessorAdmin(ImportExportModelAdmin):
    list_display = ('name', 'department')
    search_fields = ('name',)
    list_filter = ('department__name',)


class CourseAdmin(ImportExportModelAdmin):
    list_display = ('name', 'department', 'professor', 'seats', 'prerequisite')
    search_fields = ('name', 'professor__name', 'department__name')
    list_filter = ('department__name', 'professor__name', 'department__faculty__university__name')


class StudentResource(resources.ModelResource):
    user = fields.Field(
        column_name='user__username',
        attribute='user',
        widget=ForeignKeyWidget(User, 'username')
    )
    department = fields.Field(
        column_name='department__name',
        attribute='department',
        widget=ForeignKeyWidget(Department, 'name')
    )

    class Meta:
        model = Student
        fields = ('id', 'user__username', 'user__email', 'department__name', 'average_grade', 'current_course')
        import_id_fields = ['id']

    def before_import_row(self, row, **kwargs):
        # Генерація унікального ID для студента
        if not row.get('id'):
            max_id = Student.objects.aggregate(max_id=models.Max('id'))['max_id']
            row['id'] = (max_id or 0) + 1

        # Перевірка, чи відсутнє поле user__username в CSV-файлі
        if not row.get('user__username'):
            username = self.generate_unique_username()
            email = self.generate_random_email()
            user = User.objects.create_user(username=username, email=email, password='temporarypassword')
            row['user__username'] = username
            row['user__email'] = email
            row['user'] = user.id
        else:
            if not row.get('user__email'):
                email = self.generate_random_email()
                row['user__email'] = email

            try:
                user = User.objects.get(username=row['user__username'])
                row['user'] = user.id
            except User.DoesNotExist:
                email = row['user__email'] if row.get('user__email') else self.generate_random_email()
                user = User.objects.create_user(username=row['user__username'], email=email, password='temporarypassword')
                row['user'] = user.id

        # Перевірка та створення відділу, якщо він вказаний в CSV-файлі
        if row.get('department__name'):
            department, created = Department.objects.get_or_create(name=row['department__name'])
            row['department'] = department.id
        else:
            raise ValueError("Department is required for each student")

    def generate_unique_username(self):
        while True:
            username = 'user' + ''.join(random.choices(string.ascii_lowercase + string.digits, k=6))
            if not User.objects.filter(username=username).exists():
                return username

    def generate_random_email(self):
        random_string = ''.join(random.choices(string.ascii_lowercase + string.digits, k=10))
        random_digits = ''.join(random.choices(string.digits, k=3))
        email = f"{random_string}{random_digits}@gmail.com"
        return email


class StudentAdmin(ImportExportModelAdmin):
    resource_class = StudentResource
    list_display = ('user', 'department', 'average_grade', 'current_course', 'download_link')
    search_fields = ('user__username', 'department__name')
    list_filter = ('department__name', 'department__faculty__name')
    actions = ['generate_idp_single', 'generate_idp_group', 'process_applications']

    def download_link(self, obj):
        if obj.document:
            url = reverse('download_document', args=[obj.document.name])
            return format_html("<a href='{}' download='{}'>Завантажити</a>", obj.document.url.rstrip('/'),
                               obj.document.name)
        return "Файл відсутній"

    download_link.short_description = 'Документ'

    def generate_idp_single(self, request, queryset):
        for student in queryset:
            generate_idp_for_student(student.id)

    generate_idp_single.short_description = "Генерувати ІНП для вибраних студентів"

    def generate_idp_group(self, request, queryset):
        department_ids = set(queryset.values_list('department__id', flat=True))
        for department_id in department_ids:
            generate_idp_for_group(department_id)

    generate_idp_group.short_description = "Генерувати ІНП для груп студентів за кафедрами"

    def process_applications(self, request, queryset):
        output = io.StringIO()
        call_command('process_applications', stdout=output)
        self.message_user(request, output.getvalue())

    process_applications.short_description = "Запустити алгоритм багатокритеріального вибору"

    def generate_log_entries(self, result, request):
        from django.contrib.admin.models import LogEntry, ADDITION, CHANGE
        from django.contrib.contenttypes.models import ContentType

        for row_result in result.rows:
            if hasattr(row_result, 'object'):
                object_repr = str(row_result.object)
                action_flag = ADDITION if row_result.import_type == RowResult.IMPORT_TYPE_NEW else CHANGE
                LogEntry.objects.log_action(
                    user_id=request.user.pk,
                    content_type_id=ContentType.objects.get_for_model(row_result.object.__class__).pk,
                    object_id=row_result.object.pk,
                    object_repr=object_repr,
                    action_flag=action_flag,
                )
            else:
                object_repr = 'Unknown object'
                action_flag = ADDITION if row_result.import_type == RowResult.IMPORT_TYPE_NEW else CHANGE
                LogEntry.objects.log_action(
                    user_id=request.user.pk,
                    content_type_id=ContentType.objects.get_for_model(Student).pk,
                    object_id=None,
                    object_repr=object_repr,
                    action_flag=action_flag,
                )


class CourseApplicationAdmin(ImportExportModelAdmin):
    list_display = ('student', 'course', 'priority')
    search_fields = ('student__user__username', 'course__name')
    list_filter = ('course__name', 'student__department__name')


class CourseEnrollmentAdmin(ImportExportModelAdmin):
    list_display = ('student', 'course')
    search_fields = ('student__user__username', 'course__name')
    list_filter = ('course__name', 'student__department__name')
    actions = ['export_to_docx']

    def export_to_docx(self, request, queryset):
        # Create a new Document
        document = Document()

        # Group enrollments by course
        enrollments_by_course = {}
        for enrollment in queryset:
            if enrollment.course not in enrollments_by_course:
                enrollments_by_course[enrollment.course] = []
            enrollments_by_course[enrollment.course].append(enrollment.student)

        # Add a heading to the document
        document.add_heading('Звіт по сформованим групам', level=1)

        # Iterate over each course and add its enrollments to the document
        for course, students in enrollments_by_course.items():
            document.add_heading(f"{course.name} - Викладач: {course.professor.name} - Зайнятих місць: {len(students)}", level=2)
            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Студент'
            hdr_cells[1].text = 'Електронна пошта'
            hdr_cells[2].text = 'Курс'
            hdr_cells[3].text = 'Кафедра'

            for student in students:
                row_cells = table.add_row().cells
                row_cells[0].text = f"{student.user.first_name} {student.user.last_name}"
                row_cells[1].text = student.user.email
                row_cells[2].text = str(student.current_course)
                row_cells[3].text = student.department.name

            document.add_paragraph('')  # Add a blank line after each course

        # Generate the response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=course_enrollments.docx'
        document.save(response)
        return response

    export_to_docx.short_description = "Експортувати в DOCX"


class CustomUserAdmin(UserAdmin):
    change_list_template = "admin/auth/user/change_list.html"

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('bulk_add/', self.admin_site.admin_view(self.bulk_add_users), name='bulk_add_users'),
            path('export_users/', self.admin_site.admin_view(self.export_users), name='export_users'),
        ]
        return custom_urls + urls

    def bulk_add_users(self, request):
        if request.method == 'POST':
            form = BulkUserCreationForm(request.POST)
            if form.is_valid():
                users = form.save()
                for user in users:
                    temporary_passwords[user.username] = user.raw_password  # Зберігаємо незахищені паролі
                self.message_user(request, f"Створено {len(users)} користувачів.")
                return redirect('..')
        else:
            form = BulkUserCreationForm()

        context = {
            'form': form,
            'opts': self.model._meta,
            'app_label': self.model._meta.app_label,
        }
        return render(request, 'admin/bulk_add_users.html', context)

    def export_users(self, request):
        users = User.objects.all()
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="users_{slugify(str(datetime.now()))}.csv"'

        writer = csv.writer(response)
        writer.writerow(['ID', 'Username', 'Email', 'Password'])

        for user in users:
            password = temporary_passwords.get(user.username, "N/A")
            writer.writerow([user.id, user.username, user.email, password])

        return response


admin.site.unregister(User)
admin.site.register(User, CustomUserAdmin)
admin.site.register(University, UniversityAdmin)
admin.site.register(Faculty, FacultyAdmin)
admin.site.register(Department, DepartmentAdmin)
admin.site.register(Course, CourseAdmin)
admin.site.register(Student, StudentAdmin)
admin.site.register(CourseEnrollment, CourseEnrollmentAdmin)
admin.site.register(CourseApplication, CourseApplicationAdmin)
admin.site.register(Professor, ProfessorAdmin)
